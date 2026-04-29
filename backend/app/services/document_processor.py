import asyncio
import io
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader

from app.core.config import get_settings


class DocumentProcessor:
    """Extracts text, chunks content, and attaches enterprise metadata."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    def __init__(self) -> None:
        self.settings = get_settings()
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=1200,
            chunk_overlap=180,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    async def save_upload(self, upload: UploadFile) -> tuple[str, bytes]:
        filename = upload.filename or "uploaded-file"
        suffix = Path(filename).suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type '{suffix or 'unknown'}'. Supported types: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}",
            )

        file_bytes = await upload.read()
        max_bytes = self.settings.max_upload_size_mb * 1024 * 1024
        if not file_bytes:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")
        if len(file_bytes) > max_bytes:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File exceeds {self.settings.max_upload_size_mb} MB limit",
            )

        destination = self.settings.upload_dir / f"{uuid.uuid4()}_{upload.filename}"
        await asyncio.to_thread(destination.write_bytes, file_bytes)
        return str(destination), file_bytes

    async def to_documents(
        self,
        upload: UploadFile,
        allowed_roles: list[str],
        uploaded_by: str,
        document_id: str,
    ) -> tuple[list[Document], str]:
        storage_path, file_bytes = await self.save_upload(upload)
        pages = await asyncio.to_thread(self._extract_pages, upload.filename or "uploaded-file", file_bytes)

        documents: list[Document] = []
        for page_number, page_text in enumerate(pages, start=1):
            cleaned_page = page_text.strip()
            if not cleaned_page:
                continue

            for chunk_index, chunk in enumerate(self.splitter.split_text(cleaned_page)):
                if not chunk.strip():
                    continue
                metadata = {
                    "document_id": document_id,
                    "source": upload.filename,
                    "page": page_number,
                    "chunk_id": f"{document_id}-{page_number}-{chunk_index}",
                    "allowed_roles": ",".join(allowed_roles),
                    "uploaded_by": uploaded_by,
                }
                documents.append(Document(page_content=chunk, metadata=metadata))

        if not documents:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No readable text could be extracted from the uploaded file",
            )

        return documents, storage_path

    def _extract_pages(self, filename: str, file_bytes: bytes) -> list[str]:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            return [page.extract_text() or "" for page in reader.pages]
        if suffix == ".docx":
            document = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
            return [text]
        return [file_bytes.decode("utf-8", errors="ignore")]
